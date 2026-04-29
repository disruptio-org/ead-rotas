#!/usr/bin/env python3
import argparse
import csv
import re
import shutil
from collections import defaultdict
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from openpyxl import load_workbook
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


VEHICLES = {
    "lisboa": {"vehicle": "AA-14-HG", "driver": "Tiago", "assistant": "Daniel", "base_route": "Porto Alto → Lisboa → Porto Alto"},
    "santarem": {"vehicle": "AL-96-DL", "driver": "Arménio", "assistant": "—", "base_route": "Porto Alto → Santarém → Porto Alto"},
    "alentejo": {"vehicle": "37-ZH-01", "driver": "Marcus", "assistant": "Rui", "base_route": "Porto Alto → Évora → Porto Alto"},
}

LISBOA_KEYS = ["lisboa", "oeiras", "cascais", "sintra", "amadora", "almada", "seixal", "barreiro", "setubal", "setúbal", "sesimbra", "colombo", "oriente", "expo"]
SANTAREM_KEYS = ["santarem", "santarém", "tomar", "torres novas", "abrantes", "entroncamento", "cartaxo", "benavente", "rio maior"]
ALENTEJO_KEYS = ["evora", "évora", "beja", "elvas", "portalegre", "estremoz", "sines"]
ALGARVE_KEYS = ["faro", "portimao", "portimão", "albufeira", "loule", "loulé", "tavira", "lagos"]
NORTH_KEYS = ["coimbra", "aveiro", "viseu", "porto", "braga", "guimaraes", "guimarães"]

PT_MONTHS = {
    "janeiro": 1, "fevereiro": 2, "marco": 3, "março": 3, "abril": 4, "maio": 5, "junho": 6,
    "julho": 7, "agosto": 8, "setembro": 9, "outubro": 10, "novembro": 11, "dezembro": 12,
}
WEEKDAYS = {
    "segunda": 0, "segunda-feira": 0, "monday": 0,
    "terca": 1, "terça": 1, "terca-feira": 1, "terça-feira": 1, "tuesday": 1,
    "quarta": 2, "quarta-feira": 2, "wednesday": 2,
    "quinta": 3, "quinta-feira": 3, "thursday": 3,
    "sexta": 4, "sexta-feira": 4, "friday": 4,
    "sabado": 5, "sábado": 5, "saturday": 5,
    "domingo": 6, "sunday": 6,
}


@dataclass
class ServiceRow:
    service_date: date
    rsad: str
    client: str
    location: str
    service: str
    raw: Dict[str, str]


def norm(s: Optional[str]) -> str:
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s)).strip()


def norm_key(s: Optional[str]) -> str:
    s = norm(s).lower()
    repl = str.maketrans("áàãâéêíóôõúç", "aaaaeeiooouc")
    return s.translate(repl)


def load_rows(path: Path) -> List[Dict[str, str]]:
    if path.suffix.lower() == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            return [dict(r) for r in reader]
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    headers = [norm(h) for h in rows[0]]
    out = []
    for row in rows[1:]:
        if all(v is None or norm(v) == "" for v in row):
            continue
        item = {}
        for i, h in enumerate(headers):
            item[h] = row[i]
        out.append(item)
    return out


def pick_field(headers: List[str], candidates: List[str]) -> Optional[str]:
    nk = {norm_key(h): h for h in headers}
    for cand in candidates:
        if cand in nk:
            return nk[cand]
    for h in headers:
        hk = norm_key(h)
        if any(cand in hk for cand in candidates):
            return h
    return None


def parse_date_value(v) -> Optional[date]:
    if v is None or norm(v) == "":
        return None
    if isinstance(v, datetime):
        return v.date()
    if isinstance(v, date):
        return v
    s = norm(v)
    for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d", "%Y-%m-%d %H:%M:%S.%f", "%Y-%m-%d %H:%M:%S"]:
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    # excel serial fallback is intentionally omitted
    return None


def canonicalize(rows: List[Dict[str, str]]) -> List[ServiceRow]:
    if not rows:
        return []
    headers = list(rows[0].keys())
    date_field = pick_field(headers, ["data do servico", "data servico", "data", "date"])
    rsad_field = pick_field(headers, ["rsad", "nservico", "n servico", "nservico", "service id"])
    client_field = pick_field(headers, ["cliente", "cli_nome", "client"])
    location_field = pick_field(headers, ["morada / local", "morada", "localdesc", "local"])
    service_field = pick_field(headers, ["servico a realizar", "servico a executar", "eventtypedesc", "servico"])

    out = []
    for row in rows:
        d = parse_date_value(row.get(date_field)) if date_field else None
        if not d:
            continue
        out.append(ServiceRow(
            service_date=d,
            rsad=norm(row.get(rsad_field)) if rsad_field else "",
            client=norm(row.get(client_field)) if client_field else "",
            location=norm(row.get(location_field)) if location_field else "",
            service=norm(row.get(service_field)) if service_field else "",
            raw={k: norm(v) for k, v in row.items()},
        ))
    return out


def infer_anchor_date(rows: List[ServiceRow]) -> date:
    dates = sorted({r.service_date for r in rows})
    today = datetime.now().date()
    for d in dates:
        if d >= today:
            return d
    return dates[0]


def parse_timeframe(text: str, rows: List[ServiceRow]) -> Tuple[date, date]:
    t = norm_key(text)
    anchor = infer_anchor_date(rows)
    all_dates = sorted({r.service_date for r in rows})
    if not all_dates:
        raise ValueError("No valid service dates found in the spreadsheet.")

    if any(x in t for x in ["amanha", "tomorrow"]):
        d = anchor + timedelta(days=1)
        return d, d
    if any(x in t for x in ["hoje", "today", "este dia", "this day"]):
        return anchor, anchor
    if any(x in t for x in ["proxima semana", "próxima semana", "next week"]):
        monday = anchor - timedelta(days=anchor.weekday())
        if anchor.weekday() == 0:
            monday = anchor
        else:
            monday = monday + timedelta(days=7)
        return monday, monday + timedelta(days=4)
    if any(x in t for x in ["proximo mes", "próximo mês", "next month"]):
        y = anchor.year + (1 if anchor.month == 12 else 0)
        m = 1 if anchor.month == 12 else anchor.month + 1
        start = date(y, m, 1)
        end = date(y + (1 if m == 12 else 0), 1 if m == 12 else m + 1, 1) - timedelta(days=1)
        return start, end

    m = re.search(r"semana de (\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4})", t)
    if m:
        d = parse_date_value(m.group(1))
        monday = d - timedelta(days=d.weekday())
        return monday, monday + timedelta(days=4)

    m = re.search(r"(\d{1,2}/\d{1,2}/\d{4}|\d{4}-\d{2}-\d{2})", t)
    if m and " a " not in t and " to " not in t:
        d = parse_date_value(m.group(1))
        return d, d

    m = re.search(r"(\d{1,2})\s*(?:a|to)\s*(\d{1,2})\s+de\s+([a-zçãé]+)", t)
    if m:
        d1, d2, month_name = int(m.group(1)), int(m.group(2)), m.group(3)
        month = PT_MONTHS[month_name]
        year = anchor.year
        return date(year, month, d1), date(year, month, d2)

    wd_matches = [name for name in WEEKDAYS if name in t]
    if len(wd_matches) >= 2:
        wd1, wd2 = WEEKDAYS[wd_matches[0]], WEEKDAYS[wd_matches[1]]
        week_start = anchor - timedelta(days=anchor.weekday())
        start = week_start + timedelta(days=wd1)
        end = week_start + timedelta(days=wd2)
        if end < start:
            end = start
        return start, end

    raise ValueError("Could not understand the requested timeframe.")


def classify_cluster(row: ServiceRow) -> str:
    blob = norm_key(" ".join([row.client, row.location, row.service]))
    if any(k in blob for k in ALGARVE_KEYS):
        return "excluded"
    if any(k in blob for k in NORTH_KEYS):
        return "excluded"
    if any(k in blob for k in ALENTEJO_KEYS):
        return "alentejo"
    if any(k in blob for k in SANTAREM_KEYS):
        return "santarem"
    if any(k in blob for k in LISBOA_KEYS):
        return "lisboa"
    return "lisboa"


def exclusion_reason(row: ServiceRow) -> Optional[str]:
    missing = []
    if not row.rsad or row.rsad.upper() == "NULL":
        missing.append("RSAD")
    if not row.client or row.client.upper() == "NULL":
        missing.append("Cliente")
    if not row.service or row.service.upper() == "NULL":
        missing.append("Serviço")
    if missing:
        return f"Dados críticos em falta: {', '.join(missing)}"
    cluster = classify_cluster(row)
    if cluster == "excluded":
        return "Fora de eixo para planeamento standard da skill"
    return None


def assign_routes(rows: List[ServiceRow]) -> Tuple[Dict[date, Dict[str, List[ServiceRow]]], Dict[date, List[Tuple[ServiceRow, str]]]]:
    routes = defaultdict(lambda: defaultdict(list))
    excluded = defaultdict(list)
    for row in rows:
        reason = exclusion_reason(row)
        if reason:
            excluded[row.service_date].append((row, reason))
            continue
        cluster = classify_cluster(row)
        routes[row.service_date][cluster].append(row)
    return routes, excluded


def set_identification_table(doc: Document, dt: date, vehicle: str, driver: str, assistant: str):
    if not doc.tables:
        return
    table = doc.tables[0]
    vals = [dt.strftime("%d/%m/%Y"), vehicle, driver, assistant]
    for i, val in enumerate(vals):
        if i < len(table.rows) and len(table.rows[i].cells) > 1:
            table.rows[i].cells[1].text = val


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def find_paragraph(doc: Document, contains_text: str) -> Optional[Paragraph]:
    target = norm_key(contains_text)
    for p in iter_paragraphs(doc):
        if target in norm_key(p.text):
            return p
    return None


def replace_next_nonempty_paragraph(anchor: Paragraph, new_text: str):
    el = anchor._element
    nxt = el.getnext()
    while nxt is not None:
        if nxt.tag.endswith('}p'):
            para = Paragraph(nxt, anchor._parent)
            if norm(para.text):
                para.text = new_text
                return
        nxt = nxt.getnext()


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def clear_between(start_p: Paragraph, end_p: Paragraph):
    cur = start_p._element.getnext()
    while cur is not None and cur is not end_p._element:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def render_service_doc(template: Path, out_path: Path, dt: date, cluster: str, services: List[ServiceRow]):
    cfg = VEHICLES[cluster]
    doc = Document(template)
    set_identification_table(doc, dt, cfg["vehicle"], cfg["driver"], cfg["assistant"])
    seq_anchor = find_paragraph(doc, "Sequência da Rota")
    if seq_anchor:
        route_nodes = [r.location for r in services if r.location and r.location.upper() != "NULL"]
        route_string = cfg["base_route"]
        if route_nodes:
            unique_nodes = []
            for n in route_nodes:
                if n not in unique_nodes:
                    unique_nodes.append(n)
            core = " → ".join(unique_nodes[:6])
            route_string = f"Porto Alto → {core} → Porto Alto"
        replace_next_nonempty_paragraph(seq_anchor, route_string)
    plan_anchor = find_paragraph(doc, "Plano de Execução")
    checklist_anchor = find_paragraph(doc, "Checklist Operacional")
    if plan_anchor and checklist_anchor:
        clear_between(plan_anchor, checklist_anchor)
        cur = plan_anchor
        for s in services:
            cur = insert_paragraph_after(cur, f"RSAD: {s.rsad}")
            cur = insert_paragraph_after(cur, f"Cliente: {s.client}")
            cur = insert_paragraph_after(cur, f"Morada: {s.location or 'NULL'}")
            cur = insert_paragraph_after(cur, f"Serviço a realizar: {s.service}")
            cur = insert_paragraph_after(cur, "")
    doc.save(out_path)


def render_excluded_doc(template: Path, out_path: Path, dt: date, excluded: List[Tuple[ServiceRow, str]]):
    doc = Document(template)
    title_p = None
    for p in doc.paragraphs:
        if "SERVIÇOS A REMARCAR" in p.text.upper():
            title_p = p
            break
    if title_p is None:
        title_p = doc.paragraphs[0]
    # remove all paragraphs after title
    cur = title_p._element.getnext()
    while cur is not None:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt
    p = insert_paragraph_after(title_p, dt.strftime("%d/%m/%Y"))
    for s, reason in excluded:
        p = insert_paragraph_after(p, f"{s.rsad or 'SEM RSAD'} | {s.client or 'SEM CLIENTE'} | {s.location or 'SEM LOCAL'} | {s.service or 'SEM SERVIÇO'} | {reason}")
    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True)
    ap.add_argument("--timeframe", required=True)
    ap.add_argument("--service-template", required=True)
    ap.add_argument("--excluded-template", required=True)
    ap.add_argument("--output-dir", required=True)
    args = ap.parse_args()

    rows = canonicalize(load_rows(Path(args.input)))
    if not rows:
        raise SystemExit("No usable rows found.")
    start, end = parse_timeframe(args.timeframe, rows)
    selected = [r for r in rows if start <= r.service_date <= end]
    if not selected:
        raise SystemExit(f"No rows found between {start} and {end}.")

    routes, excluded = assign_routes(selected)
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    generated = []
    for dt in sorted(routes.keys() | excluded.keys()):
        for cluster, services in routes.get(dt, {}).items():
            cfg = VEHICLES[cluster]
            fname = f"Folha_Servico_{dt.isoformat()}_{cfg['vehicle']}.docx"
            out = out_dir / fname
            render_service_doc(Path(args.service_template), out, dt, cluster, services)
            generated.append(out)
        if excluded.get(dt):
            out = out_dir / f"Servicos_Excluidos_{dt.isoformat()}.docx"
            render_excluded_doc(Path(args.excluded_template), out, dt, excluded[dt])
            generated.append(out)

    print("Generated files:")
    for path in generated:
        print(path)

if __name__ == "__main__":
    main()
